from typing import Any. Optional
from datetime import date
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from tkinter import Event

from utilities import PrinterApp, Function, MIN_DATE

from os import getcwd
from os.path import isfile, join
from io import StringIO
from json import loads
from re import compile as recompile
from functools import wraps

from unidecode import unidecode
from docx import Document
from docx.shared import Pt

from tkinter import StringVar as StrVar, BooleanVar as BoolVar, Text
from tkinter.ttk import Entry, Button, Treeview, Scrollbar, Checkbutton, Radiobutton
from tkinter.messagebox import showinfo, askokcancel
from tkinter.filedialog import askdirectory
from tkcalendar import DateEntry

class TaxPrinter(PrinterApp):
    # decorators #
    def check(func: Function) -> Function:
        @wraps(func)
        def wrapper(self, *args: Any, **kwargs: Any) -> None:
            func(self, *args, **kwargs)
            state: str = self.get_state(self.elem.tree_selected.get_children())
            self.elem.btn_print.config(state=state)
            self.elem.btn_name.config(state=state)
        return wrapper

    def sort(func: Function) -> Function:
        @wraps(func)
        def wrapper(self, *args, **kwargs) -> None:
            func(self, *args, **kwargs)
            # sort by point, then by project #
            vals: list[tuple[str, str, str]] = [
              (iid, self.elem.tree_selected.set(iid, 'point'), self.elem.tree_selected.set(iid, 'project'))
              for iid in self.elem.tree_selected.get_children()
              ]
            for i, (iid, *_) in enumerate(self.sort2return(vals, key=lambda x: (x[2], self.point2tuple(x[1]), x[0]))):
                self.elem.tree_selected.move(iid, '', i)
        return wrapper

    def set_data(self) -> None:
        # check if file exists #
        if not isfile(self.vars.file):
            self.throw_error(1)
            return

        # try to read data #
        data: list[dict[str, Any]] = []
        try:
            with open(self.vars.file, 'rt', encoding='utf-8') as f:
                data.extend(loads(f.read()))

        except Exception as e:
            self.throw_error(0, str(e))

        else:
            for i, project in enumerate(data):
                self.elem.tree_all.insert('', 'end', i, text=project['name'], values=(
                  project['description'],
                  *self.flatten((dates['from'], dates['to']) for dates in project['dates'])
                  ), open=False, tags=('catalogue', ))
                for point in project['points']:
                    self.elem.tree_all.insert(i, 'end', text=point['point'], values=(point['text'], ))

    @check
    def clear_data(self) -> None:
        # clear data #
        self.elem.tree_all.delete(*self.elem.tree_all.get_children())
        self.elem.tree_selected.delete(*self.elem.tree_selected.get_children())

    def toggle(self) -> None:
        states: tuple[str, str] = self.vars.var.cash.get(), self.vars.var.addons.get()
        cash_state, addons_state = tuple(self.get_state(x) for x in states)
        self.elem.entry_cash.config(state=cash_state)
        self.elem.radbtn_auto.config(state=cash_state)
        self.elem.radbtn_all.config(state=cash_state)
        self.elem.entry_addons.config(state=addons_state)

    def set_text(self) -> None:
        self.elem.txt_opening.delete('1.0', 'end')
        self.elem.txt_opening.insert('1.0', self.vars.var[self.vars.var.opening_mode.get()].get())

    def set_path(self) -> None:
        # get path #
        if path := askdirectory(title='Wybierz folder', initialdir=self.vars.var.filepath.get()):
            self.vars.var.filepath.set(path)

    def make_name(self) -> None:
        # get values #
        name: list[str] = []
        values: tuple[tuple[str, str, str], ...] = tuple(
          self.elem.tree_selected.item(iid, 'values')[:3]
          for iid in self.elem.tree_selected.get_children()
          )

        # make new filename #
        if self.vars.var.opening_mode.get() == 'contract_text':
            name.append('u')

        # for single point #
        if len(values) == 1:
            project, point, _ = values[0]
            name.append(self.vars.patterns.single_name.sub('', unidecode(project).lower()))
            name.append(point.replace('.', ''))

        # for mutiple points #
        else:
            name.extend(self.sort2return({
              self.vars.patterns.multi_name.sub('', unidecode(project).lower())
              for project, *_ in values
              }))

        # extract dates #
        day: date = self.str2date(self.vars.var.date.get())
        dates: tuple[Optional[tuple[date, date]], ...] = tuple(
          self.extract_dates(day, self.pair_up(
            self.str2date(d)
            for d in self.elem.tree_all.item(project_iid, 'values')[1:]
            ))
          for *_, project_iid in values
          ))

        # find edge dates #
        if all(dates):
            starts, ends = zip(*dates)
            start, end = max(starts), min(ends)
            if start <= end:
                name.append(self.get_timespan_desc(start, end))

        # set the filename #
        self.vars.var.filename.set('_'.join(name))

    @sort
    @check
    def add(self) -> None:
        # get all iids and leave out folders #
        iids: tuple[tuple[str, str], ...] = tuple(
          (iid, self.elem.tree_all.parent(iid))
          for iid in self.elem.tree_all.selection()
          if not self.elem.tree_all.tag_has('catalogue', iid)
          )

        # check wherever iids correct #
        if not iids:
            self.throw_error(201)
            return

        # check wherever already not selected #
        clear_iids: tuple[str, ...] = tuple(iid for iid, _ in iids)
        if any(iid in clear_iids for iid in tuple(
          self.elem.tree_selected.item(child, 'values')[3]
          for child in self.elem.tree_selected.get_children()
          )):
            self.throw_error(602)
            return

        # add elements to table #
        for iid, parent_iid in iids:
            self.elem.tree_selected.insert('', 'end', values=(
              self.elem.tree_all.item(parent_iid, 'text'),
              self.elem.tree_all.item(iid, 'text'),
              parent_iid,
              iid
              ))

    @sort
    @check
    def add_all(self) -> None:
        # clear table #
        self.remove_all()

        # refill table #
        for parent_iid in self.elem.tree_all.get_children():
            for iid in self.elem.tree_all.get_children(parent_iid):
                self.elem.tree_selected.insert('', 'end', values=(
                  self.elem.tree_all.item(parent_iid, 'text'),
                  self.elem.tree_all.item(iid, 'text'),
                  parent_iid,
                  iid
                  ))

    @sort
    @check
    def add_by_btn(self, event: Optional[Event]=None, /) -> None:
        iid: str = self.elem.tree_all.focus()

        # check wherever iid correct #
        if not iid:
            return

        # check wherever catalogue was selected #
        if self.elem.tree_all.tag_has('catalogue', iid):
            return

        # check wherever already not selected #
        if iid in tuple(self.elem.tree_selected.item(x, 'values')[3] for x in self.elem.tree_selected.get_children()):
            return

        # add element to table #
        parent_iid: str = self.elem.tree_all.parent(iid)
        self.elem.tree_selected.insert('', 'end', values=(
          self.elem.tree_all.item(parent_iid, 'text'),
          self.elem.tree_all.item(iid, 'text'),
          parent_iid,
          iid
          ))

    @check
    def remove(self) -> None:
        iids: list[str] = self.elem.tree_selected.selection()

        # check wherever iids correct #
        if not iids:
            self.throw_error(201)
            return

        # remove selected elements #
        self.elem.tree_selected.delete(*iids)

    @check
    def remove_all(self) -> None:
        # remove all elements #
        self.elem.tree_selected.delete(*self.elem.tree_selected.get_children())

    @check
    def remove_by_btn(self, event: Optional[Event]=None, /) -> None:
        iid: str = self.elem.tree_selected.focus()

        # check wherever iids correct #
        if not iid:
            return

        # remove selected elements #
        self.elem.tree_selected.delete(iid)

    def print(self) -> None:
        name: str = self.vars.var.filename.get()

        # check if filename not empty #
        if not name:
            self.throw_error(3)
            return

        # check if filename correct #
        if any(char in name for char in self.vars.bad_chars):
            self.throw_error(2)
            return

        path: str = join(self.vars.var.filepath.get(), f'{name}.docx')

        # check wherever file exists #
        if isfile(path):
            if not askokcancel(title='Plik już istnieje', message='Czy chcesz kontynuować?'):
                self.throw_error(5)
                return

        # get projects, and prepare text #
        txt: str = ''
        with StringIO('', newline='') as txt_file:
            if beg := self.elem.txt_opening.get('1.0', 'end-1c').strip():
                txt_file.write(beg + '<br><br>')

            items: tuple[str, ...] = self.elem.tree_selected.get_children()
            for project, parent_iid in self.sort2return({
              self.elem.tree_selected.item(iid, 'values')[::2]
              for iid in items
              }, key=lambda x: x[0]):

                # perpare variables #
                desc, *str_dates = self.elem.tree_all.item(parent_iid, 'values')
                day, *dates = tuple(self.str2date(d) for d in (self.vars.var.date.get(), *str_dates))
                time: tuple[date, date] = self.extract_dates(day, self.pair_up(dates))

                # check wherever time was found #
                if not time:
                    self.throw_error(603)
                    return

                # write text #
                txt_file.write(self.vars.var.title_text.get().strip().replace('<t>', project) + '<br>')
                txt_file.write(desc.replace('<d>', self.connect_dates(*time)) + '<br>')

                # get data for point, and write them #
                collected_points: tuple[tuple[str, str], ...] = tuple(
                  self.elem.tree_selected.item(iid, 'values')[1::2]
                  for iid in items
                  if self.elem.tree_selected.set(iid, 'project') == project
                  )

                # check for cash print requirements #
                print_cash: bool = \
                  (self.vars.var.cash_mode.get() == 'auto' and \
                  1 < len(collected_points)) or \
                  self.vars.var.cash_mode.get() == 'all'

                for point, iid in collected_points:
                    if self.vars.var.cash.get() and print_cash:
                        txt_file.write(self.vars.var.cash_text.get())

                    txt_file.write(self.replace_mutiple(self.vars.var.point_text.get(), {
                      '<p>': point,
                      '<o>': self.elem.tree_all.item(iid, 'values')[0]
                      }))

                    if self.vars.var.addons.get():
                        txt_file.write(self.vars.var.addons_text.get())

                    txt_file.write('<br>')
                txt_file.write('<br>')

            txt = txt_file.getvalue().replace('<br>', '\n').strip()
        
        # create file #
        document: Document = Document()

        # define style #
        style: ParagraphStyle = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # regex matching #
        tags: list[str] = self.vars.patterns.tags.findall(txt)
        txt: list[str] = self.vars.patterns.tags.split(txt)

        # write data to docx #
        par: Paragraph = document.add_paragraph(txt.pop(0), style)
        for tag, part in zip(tags, txt):
            run: Run = par.add_run(part)

            # set formatting #
            match tag:
                case '<b>':
                    run.font.bold = True
                case '</b>':
                    run.font.bold = False
                case '<i>':
                    run.font.italic = True
                case '</i>':
                    run.font.italic = False
                case '<u>':
                    run.font.underline = True
                case '</u>':
                    run.font.underline = False
                case '<s>':
                    run.font.strike = True
                case '</s>':
                    run.font.strike = False

        # save file #
        try:
            document.save(path)

        except Exception as e:
            self.throw_error(0, str(e))

        else:
            showinfo(title='Zapisywanie', message='Sukces')

    def reload(self) -> None:
        # reload data #
        self.clear_data()
        self.set_data()

    # overridden #
    def pre(self) -> None:
        # important #
        self.vars.title = 'Tax Printer'
        self.vars.size.update(
          min = (400, 450),
          max = (650, 650)
          )

        # add, and set variables #
        self.vars.update(
          tags = ('b', 'i', 'u', 's'),
          file = join(getcwd(), 'data.json')
          )
        self.vars.var.update(
          date = StrVar(),
          filename = StrVar(value='opis'),
          filepath = StrVar(value=getcwd()),
          title_text = StrVar(value='<b><t></b>'),
          point_text = StrVar(value='<b>Punkt <p></b> „<o>”'),
          cash = BoolVar(),
          addons = BoolVar(),
          cash_text = StrVar(value='Suma <b>............... PLN</b>, '),
          addons_text = StrVar(value=' - '),
          cash_mode = StrVar(),
          opening_mode = StrVar(),
          facture_text = StrVar(value=(
            'Akapit przykładowy 1.\n'
            'Tutaj wpisać tekst na, który ma pojawić się na początku dokumentu'
            )),
          contract_text = StrVar(value=(
            'Akapit przykładowy 2.\n'
            'Tutaj wpisać tekst na, który ma pojawić się na początku dokumentu'
            ))
          )
        self.vars.patterns.update(
          single_name = recompile(r'[^a-z0-9]+'),
          multi_name = recompile(r'[^a-z]+'),
          tags = recompile(r'</?(?:{})>'.format('|'.join(self.vars.tags)))
          )
        self.elem.update(
          tree_all = {
            'type': Treeview,
            'args': {'selectmode': 'extended', 'show': 'tree'},
            'grid': {'row': 0, 'column': 0, 'rowspan': 5, 'columnspan': 2},
            'sticky': 'NWES'
            },
          scroll_all = {
            'type': Scrollbar,
            'args': {'orient': 'vertical'},
            'grid': {'row': 0, 'column': 1, 'rowspan': 5},
            'sticky': 'NES'
            },
          btn_add = {
            'type': Button,
            'args': {'text': '\u25b6', 'command': self.add},
            'grid': {'row': 0, 'column': 2},
            'tooltip': 'Dodaj element(y)'
            },
          btn_add_all = {
            'type': Button,
            'args': {'text': '\u25b6\u25b6', 'command': self.add_all},
            'grid': {'row': 1, 'column': 2},
            'tooltip': 'Dodaj wszystkie elementy'
            },
          btn_remove = {
            'type': Button,
            'args': {'text': '\u25c0', 'command': self.remove},
            'grid': {'row': 3, 'column': 2},
            'tooltip': 'Usuń element(y)'
            },
          btn_remove_all = {
            'type': Button,
            'args': {'text': '\u25c0\u25c0', 'command': self.remove_all},
            'grid': {'row': 4, 'column': 2},
            'tooltip': 'Usuń wszystkie elementy'
            },
          tree_selected = {
            'type': Treeview,
            'args': {'selectmode': 'extended', 'show': 'headings', 'columns': ('project', 'point')},
            'grid': {'row': 0, 'column': 3, 'rowspan': 5, 'columnspan': 2},
            'sticky': 'NWES'
            },
          txt_opening = {
            'type': Text,
            'args': {'wrap': 'word'},
            'grid': {'row': 5, 'column': 0, 'columnspan': 5},
            'sticky': 'NWES',
            'borderfull': {'highlightthickness': 1, 'highlightbackground': 'gray'}
            },
          radbtn_facture = {
            'type': Radiobutton,
            'args': {'text': 'Faktura', 'variable': self.vars.var.opening_mode, 'command': self.set_text, 'value': 'facture_text'},
            'grid': {'row': 6, 'column': 3},
            'sticky': 'W'
            },
          radbtn_contract = {
            'type': Radiobutton,
            'args': {'text': 'Umowa', 'variable': self.vars.var.opening_mode, 'command': self.set_text, 'value': 'contract_text'},
            'grid': {'row': 6, 'column': 4},
            'sticky': 'W'
            },
          entry_point = {
            'type': Entry,
            'args': {'textvariable': self.vars.var.point_text},
            'grid': {'row': 6, 'column': 0, 'columnspan': 2},
            'tooltip': 'Szablon podpunktu',
            'sticky': 'NWES'
            },
          chkbtn_cash = {
            'type': Checkbutton,
            'args': {'text': 'Dodaj pole kwoty', 'variable': self.vars.var.cash, 'command': self.toggle},
            'grid': {'row': 7, 'column': 0},
            'sticky': 'W'
            },
          entry_cash = {
            'type': Entry,
            'args': {'textvariable': self.vars.var.cash_text, 'state': 'disabled'},
            'grid': {'row': 7, 'column': 1},
            'tooltip': 'Tekst kwoty',
            'sticky': 'NWES'
            },
          radbtn_auto = {
            'type': Radiobutton,
            'args': {'text': 'Auto', 'variable': self.vars.var.cash_mode, 'value': 'auto'},
            'grid': {'row': 8, 'column': 0},
            'sticky': 'W'
            },
          radbtn_all = {
            'type': Radiobutton,
            'args': {'text': 'Wszystko', 'variable': self.vars.var.cash_mode, 'value': 'all'},
            'grid': {'row': 8, 'column': 1},
            'sticky': 'W'
            },
          chkbtn_addons = {
            'type': Checkbutton,
            'args': {'text': 'Dodaj myślniki', 'variable': self.vars.var.addons, 'command': self.toggle},
            'grid': {'row': 9, 'column': 0},
            'sticky': 'W'
            },
          entry_addons = {
            'type': Entry,
            'args': {'textvariable': self.vars.var.addons_text, 'state': 'disabled'},
            'grid': {'row': 9, 'column': 1},
            'tooltip': 'Tekst notatki',
            'sticky': 'NWES'
            },
          cal_select = {
            'type': DateEntry,
            'args': {'date_pattern': 'd.M.y', 'locale': 'pl_PL', 'mindate': MIN_DATE, 'textvariable': self.vars.var.date},
            'grid': {'row': 7, 'column': 3},
            'tooltip': 'Data wystawienia opisu'
            },
          btn_name = {
            'type': Button,
            'args': {'text': '\u270e', 'state': 'disabled', 'command': self.make_name},
            'grid': {'row': 7, 'column': 4},
            'tooltip': 'Wygeneruj szybką nazwę podpunktu(ów)'
            },
          entry_filename = {
            'type': Entry,
            'args': {'textvariable': self.vars.var.filename},
            'grid': {'row': 8, 'column': 3, 'columnspan': 2},
            'tooltip': 'Nazwa pliku',
            'sticky': 'NWES'
            },
          btn_filepath = {
            'type': Button,
            'args': {'text': '\u2026', 'command': self.set_path},
            'grid': {'row': 9, 'column': 3},
            'tooltip': 'Wybierz ścieżkę docelową',
            'sticky': 'W'
            },
          btn_print = {
            'type': Button,
            'args': {'text': '\ud83d\uddb6', 'state': 'disabled', 'command': self.print},
            'grid': {'row': 9, 'column': 4},
            'tooltip': 'Drukuj opis'
            }
          )
        self.grid.row.update({
          5: {'minsize': 80}
          })
        # self.grid.col.update()
        self.menus.update(
          menu_main = {
            'main': True,
            'elements': [
              ('menu', {'label': 'Plik', 'menu': 'menu_file'}),
              ('command', {'label': 'Formatowanie', 'command': self.show_format}),
              ('command', {'label': 'Pomoc', 'command': self.show_help})
              ]
            },
          menu_file = {
            'elements': [
              ('command', {'label': 'Wybierz\u2026', 'command': self.select_file}),
              ('command', {'label': 'Przeładuj', 'command': self.reload}),
              ('separator', None),
              ('command', {'label': 'Wyjdź', 'command': self.close})
              ]
            }
          )
        self.binds.extend([
          ('tree_all', ('<Return>', self.add_by_btn)),
          ('tree_all', ('<Right>', self.add_by_btn)),
          ('tree_all', ('<Double-Button-1>', self.add_by_btn)),
          ('tree_selected', ('<Return>', self.remove_by_btn)),
          ('tree_selected', ('<Left>', self.remove_by_btn)),
          ('tree_selected', ('<Double-Button-1>', self.remove_by_btn))
          ])

    def post(self) -> None:
        # styles, other settings, and actions #
        self.vars.style.configure('TFrame', background='white')
        self.vars.style.configure('TCheckbutton', background='white')
        self.vars.style.configure('TRadiobutton', background='white')
        self.elem.tree_selected.heading('project', text='Projekt')
        self.elem.tree_selected.heading('point', text='Punkt')
        self.elem.tree_selected.column('project', width=40, minwidth=40)
        self.elem.tree_selected.column('point', width=40, minwidth=40)
        self.elem.tree_all.configure(yscrollcommand=self.elem.scroll_all.set)
        self.elem.scroll_all.configure(command=self.elem.tree_all.yview)
        self.elem.radbtn_facture.invoke()
        self.elem.chkbtn_cash.invoke()
        self.elem.radbtn_auto.invoke()

        self.set_data()

    def close(self) -> None:
        self.root.destroy()


if __name__ == '__main__':
    TaxPrinter()
